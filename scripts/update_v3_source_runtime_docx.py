from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "解码阅读-V3阶段开发记录-2026-08-11.docx"
DOCX_FALLBACK_PATH = ROOT / "docs" / "DEVELOPMENT_RECORD_2026-08-11-stage6.docx"
SECTION_TITLE = "10. 书源本地优先运行时（2026-08-11 追加）"
STAGE2_SECTION_TITLE = "11. 书源运行时第二轮完善（2026-08-11）"
STAGE3_SECTION_TITLE = "12. YCK 全目录导入与 Android 大容量存储（2026-08-11）"
STAGE4_SECTION_TITLE = "13. 第四轮真实阅读与 Android 全量导入（2026-08-12）"
STAGE5_SECTION_TITLE = "14. 第五轮发现页运行态隔离与视频问题修复（2026-08-12）"
STAGE6_SECTION_TITLE = "15. 第六轮自适应搜索、发现跨源回退与运行池（2026-08-13）"
STAGE6_FINAL_TITLE = "15.6 最终 Android 真机复测与按钮热修（2026-08-13）"
STAGE7_SECTION_TITLE = "16. 第七阶段真实书源质量、可读率与稳定交付（2026-08-13）"
STAGE7_REPLAY_TITLE = "16.6 首窗口高频规则差异重放（2026-08-13）"
STAGE7_REPLAY2_TITLE = "16.7 第二批规则重放与外部变化审计（2026-08-13）"


def set_east_asia_font(run, name="微软雅黑"):
    run.font.name = name
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    return paragraph


def add_body(document, text, bold_prefix=""):
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        prefix.bold = True
        set_east_asia_font(prefix)
        remainder = paragraph.add_run(text[len(bold_prefix):])
        set_east_asia_font(remainder)
    else:
        run = paragraph.add_run(text)
        set_east_asia_font(run)
    return paragraph


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def style_table(table):
    table.style = "Table Grid"
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_east_asia_font(run)
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True


def save_document(document):
    try:
        document.save(DOCX_PATH)
        return DOCX_PATH
    except PermissionError:
        document.save(DOCX_FALLBACK_PATH)
        return DOCX_FALLBACK_PATH


def append_stage2_section(document):
    if any(paragraph.text.strip() == STAGE2_SECTION_TITLE for paragraph in document.paragraphs):
        return False

    document.add_page_break()
    heading = document.add_heading(STAGE2_SECTION_TITLE, level=1)
    for run in heading.runs:
        set_east_asia_font(run)

    document.add_heading("11.1 稳定失败分类与脱敏诊断", level=2)
    add_bullet(document, "新增统一 SourceRuntimeError 与 classifySourceFailure；搜索、详情、目录、正文和传输层使用稳定错误码，不再只返回笼统的请求失败。")
    add_bullet(document, "区分网络、超时、HTTP 拦截、站点失效、登录/Cookie/WebView/验证码需求、规则或解析为空、安全脚本拒绝和预算超限。")
    add_bullet(document, "健康记录增加失败阶段、HTTP 状态、可重试标记和脱敏诊断；不保存正文、Cookie、Token 或完整响应。")

    document.add_heading("11.2 高频 3.x 兼容补齐", level=2)
    add_bullet(document, "安全脚本解释器支持受控变量声明、字符串拼接、对象/数组字面量、JSON.parse/stringify 和 String()，可执行常见动态 URL + POST 请求规则。")
    add_bullet(document, "请求层仅允许 GET/POST，支持 header/body/charset；增加可取消超时、响应头 charset 识别和 GBK/GB2312 解码。")
    add_bullet(document, "规则引擎增加 JSONPath 过滤、@children 和标签链式属性，并修复嵌套同名 HTML 标签提前截断。")
    add_bullet(document, "YCK 7655 已完成动态 POST 搜索、详情、1914 章目录和正文；7628 曾完整通过，但固定复测触发验证码，按 CAPTCHA_REQUIRED 记录。")

    document.add_heading("11.3 第二轮 200 源基准", level=2)
    rows = [
        ("有效文字 JSON / 可导入", "200 / 200", "导入率 100%"),
        ("静态状态", "ready 82 / partial 30", "needs_login 10 / blocked 78"),
        ("静态完整候选", "36", "运行时外部排除 32"),
        ("运行时合格分母", "4", "完整通过 1，阅读率 25%"),
        ("引擎侧未通过", "PARSE_EMPTY 1", "SEARCH_EMPTY 2"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(["指标", "结果", "说明"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)
    add_body(document, "外部排除：NETWORK_ERROR 12、HTTP_BLOCKED 6、SITE_UNREACHABLE 4、TIMEOUT 4、HTTP_NOT_FOUND 3、HTTP_SERVER_ERROR 2、CAPTCHA_REQUIRED 1。")
    add_body(document, "38 个失败配置审计：POST 19、请求 options 21、Cookie 16、自定义 headers 9、JS 1；公开配置下载失败 0。")
    warning = document.add_paragraph()
    warning_run = warning.add_run("验收结论：25% 仍低于 ≥80% 发布目标，PR 保持草稿，不能宣称 YCK 绝大书源已可稳定阅读。")
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    warning_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    set_east_asia_font(warning_run)

    document.add_heading("11.4 自动验证、产物与下一步", level=2)
    add_bullet(document, "前端全量 95 passed；后端 SQLite 125 passed；语法检查与 git diff --check 通过。")
    add_bullet(document, "H5 生产构建成功；APK release/android-v2/V2.apk 为 1,504,334 字节，SHA-256 为 B7C810ACA13F12FA8978B79B4AA06E2862723F41A47B0FA2E0AB4CE80601050E，v1/v2/v3 签名通过。")
    add_bullet(document, "本轮仍无 Android 真机；关闭电脑后端后的五入口同源导入、覆盖安装、重启续读和断网缓存仍是发布阻断项。")
    add_number(document, "继续为 PARSE_EMPTY 和 SEARCH_EMPTY 样本补齐安全规则夹具，不扩大任意脚本权限。")
    add_number(document, "在仅手机联网且关闭 8765 的 Android 真机完成完整阅读闭环并保存脱敏证据。")
    add_number(document, "扩大运行时有效分母并跨两个时间窗口复测；达到 ≥80% 且 CI 全绿后再申请合并。")
    return True


def append_stage3_section(document):
    if any(paragraph.text.strip() == STAGE3_SECTION_TITLE for paragraph in document.paragraphs):
        return False

    document.add_page_break()
    heading = document.add_heading(STAGE3_SECTION_TITLE, level=1)
    for run in heading.runs:
        set_east_asia_font(run)

    document.add_heading("12.1 全目录入口与批量下载", level=2)
    add_bullet(document, "书源市场适配 YCK 当前 keys 搜索参数和全部筛选条件，同时解析总数、页码、每页数量和总页数。")
    add_bullet(document, "每页优先调用批量 JSON 地址 /yuedu/shuyuan/jsons；漏项时自动以 4 路并发回退到单源 JSON，不因一个失效 ID 丢弃整页。")
    add_bullet(document, "全量导入显示页码、下载、缺失、新增和覆盖进度，可保存进度后停止，并按筛选条件断点续传。")
    add_bullet(document, "合法来源全部保存；仅 ready 自动启用，登录、验证码、WebView、非文字类型等受限来源默认禁用且显示原因。")

    document.add_heading("12.2 唯一身份与 Android 分块存储", level=2)
    add_bullet(document, "新书源 ID 与 sourceKey 均使用规范化名称 + 基础 URL；同站点不同名称不再互相覆盖，旧数据 ID 保持不变。")
    add_bullet(document, "同批重复 sourceKey 复用同一 ID，后一条按覆盖策略更新，最终落盘按 ID 唯一化。")
    add_bullet(document, "新增 NovelReaderSourceStorage 原生桥，将书源按 25 条分片写入应用私有文件，通过新分片 → 新清单 → 清理旧分片完成事务式切换。")
    add_bullet(document, "本地存储架构版本提升至 4；H5 默认阻止超过 500 条的一键落盘，避免 localStorage 配额造成半写入。")

    document.add_heading("12.3 YCK 5621 条全量导入基准", level=2)
    rows = [
        ("目录条目 / 下载", "5621 / 5621", "缺失 0，无效 JSON 0"),
        ("唯一安装书源", "5327", "sourceKey 重复残留 0"),
        ("新增 / 覆盖", "5327 / 294", "合计 5621"),
        ("静态状态", "ready 2537 / partial 700", "blocked 1870 / needs_login 220"),
        ("启用 / 禁用", "2356 / 2971", "仅合格且原配置允许的来源启用"),
        ("原生存储", "214 分片", "39,669,743 字节"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(["指标", "结果", "说明"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)
    add_body(document, "脱敏报告：docs/source-acceptance/yck-full-import-stage3-2026-08-11.json；不保存正文、Cookie、Token 或完整书源 JSON。")
    warning = document.add_paragraph()
    warning_run = warning.add_run("口径说明：100% 是合法 JSON 进入导入流水线的比例，不代表全部来源可稳定阅读；完整阅读率仍未达到 ≥80% 发布门槛。")
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    warning_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    set_east_asia_font(warning_run)

    document.add_heading("12.4 自动验证、构建与下一步", level=2)
    add_bullet(document, "前端 Node 测试 96 passed；后端 SQLite 125 passed；H5 生产构建成功。")
    add_bullet(document, "APK release/android-v2/V2.apk 为 1,516,622 字节，SHA-256 为 B11EBB3669C8D4346639F31712F07A2D2882E19390D58904255CDB10FA7DD586，v1/v2/v3 签名通过。")
    add_bullet(document, "REA-AN00 已覆盖安装并成功启动 1.0.0 (10000)；本轮未直接在用户手机触发约 39.7 MB 的全量导入压力操作。")
    add_number(document, "继续处理 PARSE_EMPTY、SEARCH_EMPTY 和受控 WebView/请求脚本差异，提高真实阅读通过率。")
    add_number(document, "在关闭 8765 的 Android 真机完成全量导入耗时、磁盘占用、重启加载和随机阅读压力测试。")
    add_number(document, "漫画、音频、任意 Java 类、复杂动态加密、验证码绕过和付费内容继续保持明确边界。")
    return True


def append_stage4_section(document):
    if any(paragraph.text.strip() == STAGE4_SECTION_TITLE for paragraph in document.paragraphs):
        return False

    document.add_page_break()
    heading = document.add_heading(STAGE4_SECTION_TITLE, level=1)
    for run in heading.runs:
        set_east_asia_font(run)

    document.add_heading("13.1 规则运行时与书源市场修复", level=2)
    add_bullet(document, "兼容相对路径 POST 搜索与首页跨域跳转：先以无敏感信息 GET 确认最终站点，再向新源站重建只包含搜索参数的 POST；不跨域透传 Cookie、Authorization 或 Proxy-Authorization。")
    add_bullet(document, "完整阅读流程按斗破苍穹、剑来、诡秘之主依次尝试；只对稳定空结果码切换关键词，网络、登录、验证码和安全拒绝保持原错误。")
    add_bullet(document, "YCK 市场请求改用 APK 原生网络桥、浏览器兼容请求头和 30 秒超时；批量接口失败时自动回退为 4 路单源下载。")
    add_bullet(document, "代表源 7596 已由 PARSE_EMPTY 修复为完整通过：斗破苍穹、1914 章目录、抽样正文 1284 个清洗后字符。")

    document.add_heading("13.2 第四轮固定 200 源基准", level=2)
    rows = [
        ("有效文字 JSON / 可导入", "200 / 200", "导入率 100%，达到 ≥95%"),
        ("静态状态", "ready 83 / partial 29", "needs_login 10 / blocked 78"),
        ("静态候选 / 外部排除", "34 / 30", "外部状态均有稳定错误码"),
        ("严格运行时分母", "4", "排除无法公开访问的外部状态"),
        ("完整阅读通过", "2 / 4", "50%，仍低于 ≥80%"),
        ("引擎侧未通过", "SEARCH_EMPTY 2", "三个验收关键词均返回空结果"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(["指标", "结果", "说明"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)
    warning = document.add_paragraph()
    warning_run = warning.add_run("验收结论：本分支已达到全目录可导入和稳定分类目标，但真实完整阅读率为 50%，不能宣称 YCK 全部或绝大来源都可稳定阅读。")
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    warning_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    set_east_asia_font(warning_run)

    document.add_heading("13.3 无 8765 后端的 Android 真机全量验收", level=2)
    rows = [
        ("设备 / 系统", "REA-AN00 / Android 15", "应用 1.0.0 (10000)"),
        ("电脑后端", "8765 关闭", "无 tcp:8765 反向映射"),
        ("目录处理", "57 / 57 页", "5624 / 5624 条，缺失 0"),
        ("新增 / 覆盖", "5330 / 294", "重启后仍显示 5330 源"),
        ("桌面独立基准", "5602 / 5624 下载", "瞬时缺失 22，下载率 99.61%"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(["指标", "结果", "说明"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)
    add_body(document, "真机批量连接提前关闭时自动回退单源下载；DNS 瞬时失败时保存断点，恢复后续传。强制停止并重启后书源仍在，证明 APK 可在仅手机联网条件下完成全量导入和大容量持久化。")

    document.add_heading("13.4 自动验证、产物与后续边界", level=2)
    add_bullet(document, "前端 Node 全量 97 passed；后端 SQLite 125 passed；H5 生产构建成功。")
    add_bullet(document, "APK release/android-v2/V2.apk 为 1,520,718 字节，SHA-256 为 3C8BA0EFD6BCFB781188CA28063612CB74A1EFEACAA64B7D23D0DE656B542967，v1/v2/v3 签名通过；最终包覆盖安装并重启后仍显示 5330 源。")
    add_bullet(document, "脱敏报告保存样本 ID、哈希、状态、耗时与错误码；不保存正文、Cookie、Token 或完整响应。")
    add_number(document, "继续诊断两个 SEARCH_EMPTY 固定样本，并扩大严格运行时有效分母，在第二时间窗口复测。")
    add_number(document, "补齐 URL、文件、二维码和 3.x 深链同源去重，以及加入书架、重启续读、断网缓存的真机证据。")
    add_number(document, "Android 书源保持本地优先；后端继续只承担账号同步、云书架、云 TTS 和 H5 鉴权代理。")
    return True


def append_stage5_section(document):
    if any(paragraph.text.strip() == STAGE5_SECTION_TITLE for paragraph in document.paragraphs):
        return False

    document.add_page_break()
    heading = document.add_heading(STAGE5_SECTION_TITLE, level=1)
    for run in heading.runs:
        set_east_asia_font(run)

    document.add_heading("14.1 视频复现与根因", level=2)
    add_bullet(document, "录屏确认 5330 条书源已导入并持久化；故障发生在发现入口请求阶段，不是导入失败。")
    add_bullet(document, "阅书小说网请求 m.yueshu.org 时返回 Android DNS 异常，属于目标域名不可达，不代表手机断网或必须连接 FastAPI。")
    add_bullet(document, "原发现聚合只检查启用和规则兼容，未使用发现阶段运行结果；21853 是入口条数，不等于 21853 个入口均已验证可用。")

    document.add_heading("14.2 运行态隔离与中文错误", level=2)
    add_bullet(document, "新增独立 exploreTest，记录发现阶段 passed/failed、入口名、结果数、稳定错误码、HTTP 状态和可重试标记；与搜索 lastTest 分离。")
    add_bullet(document, "发现请求失败后保存稳定分类并暂时隔离整个来源；以后单源入口成功访问时自动恢复。")
    add_bullet(document, "Android DNS 文本统一分类为 SITE_UNREACHABLE；HTTP 404/403/429/5xx 与 DNS 均显示中文说明，不再直出英文底层异常。")
    add_bullet(document, "书源列表分别显示规则兼容性和站点运行状态，不再显示裸 compatible。")

    document.add_heading("14.3 REA-AN00 真机证据", level=2)
    rows = [
        ("初始发现聚合", "1373 源 / 21853 入口", "覆盖安装保留 5330 条书源"),
        ("DNS 失败后", "1372 源 / 21845 入口", "隔离 1 源 / 8 入口"),
        ("HTTP 404 抽样后", "1371 源 / 21836 入口", "再隔离 1 源 / 9 入口"),
        ("后端与 AI 语音", "8765 healthy", "保留 tcp:8765 ADB 映射"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(["验收项", "结果", "说明"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)
    add_body(document, "发现页口径改为同时显示来源数量与入口数量；失败卡片提供“查看其他可用入口”，不继续无意义重试已隔离 URL。")

    document.add_heading("14.4 自动验证、构建与边界", level=2)
    add_bullet(document, "前端 Node 全量 97 passed；新增 DNS 覆盖、失败隔离、成功恢复、中文提示测试，最终定向回归与 git diff --check 通过。")
    add_bullet(document, "H5 生产构建成功；最终 APK 为 1,520,718 字节，SHA-256 为 E882F28DA3ACFB134862081E129096FD8515A5C6A3719CCC41B865CFFE987C35，v1/v2/v3 签名及覆盖安装通过。")
    add_bullet(document, "Android 书源继续走本地原生网络桥；8765 只承担账号、同步、云 TTS 和 H5 代理，可启用但不是阅读前置条件。")
    add_number(document, "增加隔离冷却时间和批量重新检测，区分瞬时失败与长期失效。")
    add_number(document, "分层抽样 1371 个发现源，统计 DNS、HTTP、解析为空与成功占比，优先修复高频可复现差异。")
    add_number(document, "严格完整阅读率达到 ≥80% 前，不宣称 YCK 全部或绝大来源都可稳定阅读。")
    return True


def append_stage6_section(document):
    if any(paragraph.text.strip() == STAGE6_SECTION_TITLE for paragraph in document.paragraphs):
        return False

    document.add_page_break()
    heading = document.add_heading(STAGE6_SECTION_TITLE, level=1)
    for run in heading.runs:
        set_east_asia_font(run)

    document.add_heading("15.1 统一运行状态与搜索池", level=2)
    add_bullet(document, "新增兼容式 runtimeV2，分别记录 search、explore、detail、toc、content 的 untested、probing、passed、cooldown、blocked 状态，以及时间、耗时、结果数、HTTP 状态、稳定错误码、连续失败和冷却截止时间。")
    add_bullet(document, "配置指纹变化时自动清除旧失败状态；失败来源保留且不修改用户启用开关，只在对应阶段暂时隔离。旧 lastTest、exploreTest 和 health 继续兼容写入。")
    add_bullet(document, "未测试但静态支持搜索的文字源可以自动进入候选池；登录、验证码、非文字类型和安全规则越界来源不会自动请求。网络、DNS、HTTP 和解析为空分别使用分级冷却。")

    document.add_heading("15.2 自适应本地搜索、Wi-Fi 预热与发现回退", level=2)
    add_bullet(document, "Android 始终先搜索手机本地来源；后端登录结果只并行合并，后端失败或离线不能中断本地结果。默认每轮最多 20 源、并发 4、单源超时 6 秒、整轮上限 20 秒，并支持继续检测下一批。")
    add_bullet(document, "搜索结果按标题和作者去重并保留备用线路；搜索、详情、目录和正文分别记录真实运行结果，后续阶段失败时可以切换同书备用来源。")
    add_bullet(document, "应用前台且 Wi-Fi 空闲时每会话最多预热 20 个来源、并发 2；移动数据、应用隐藏、正在搜索或阅读时暂停，不增加常驻后台服务。")
    add_bullet(document, "发现标签统一归类，每个分类保留多个提供者；聚合入口最多尝试 3 个来源、并发最多 2 个，全部失败后显示稳定错误汇总，不再无限加载。")

    document.add_heading("15.3 自动测试、真实基准与真机证据", level=2)
    rows = [
        ("前端全量测试", "101 / 101", "runtimeV2、搜索、预热、发现回退均覆盖"),
        ("后端 SQLite 测试", "125 / 125", "账号、同步、代理与 AI TTS 无回归"),
        ("合法文字源导入", "200 / 200", "导入率 100%"),
        ("静态候选 / 流程测试", "36 / 36", "失败均有稳定错误分类"),
        ("严格运行时分母", "3", "未达到至少 20 个目标"),
        ("完整阅读通过", "1 / 3", "33.33%，未达到 ≥80%"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(["验收项", "结果", "说明"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)
    add_body(document, "第六轮外部排除 33 个：NETWORK_ERROR 13、TIMEOUT 6、HTTP_BLOCKED 6、SITE_UNREACHABLE 4、HTTP_NOT_FOUND 2、HTTP_SERVER_ERROR 2；严格分母内另有 SEARCH_EMPTY 2。")
    add_bullet(document, "REA-AN00 保留 5330 个来源。大容量候选池优化后，发现页框架约 4 秒可交互、完整目录约 10 秒出现；失效分类会跨多个提供者回退并明确结束。")
    add_bullet(document, "真机搜索斗破苍穹在 20 秒内探测 20 个来源但未得到结果；随后增加真实通过源冷启动排序并修复无结果页入口，代码已通过自动测试并进入最终 APK，尚待设备重新连接后复测。")
    add_bullet(document, "FastAPI 健康检查返回 ok，后端继续支持账号、同步和 AI 语音；Android 书源请求仍通过本地运行时执行。")

    document.add_heading("15.4 交付状态与发布边界", level=2)
    warning = document.add_paragraph()
    warning_run = warning.add_run("验收结论：严格分母 3、完整阅读 33.33%，尚未达到至少 20 个分母和 ≥80% 发布目标。PR #1 必须继续保持草稿，不宣称 YCK 全部或绝大来源已稳定可读。")
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    warning_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    set_east_asia_font(warning_run)
    add_bullet(document, "2026-08-13 最终 H5/APK 重建成功；APK 为 6,260,746 字节，SHA-256 为 AEBDCC9D87F7E4CA2A3B022DBAC20371848D83B1F94B0295867727D6E5B3405A，v1/v2/v3 签名通过。无线 ADB 在覆盖安装前断开，最终包尚未完成 REA-AN00 复测。")
    add_number(document, "设备重新连接后覆盖安装最终 APK，并复测搜索、发现、备用线路、断网缓存与 AI 语音试听。")
    add_number(document, "扩大第二时间窗口严格分母，优先修复 SEARCH_EMPTY 与高频可复现规则差异；继续禁止任意 eval、Java 类、文件系统、验证码绕过和付费内容访问。")
    return True


def append_stage6_final_section(document):
    if any(paragraph.text.strip() == STAGE6_FINAL_TITLE for paragraph in document.paragraphs):
        return False
    heading = document.add_heading(STAGE6_FINAL_TITLE, level=2)
    for run in heading.runs:
        set_east_asia_font(run)
    add_bullet(document, "REA-AN00 使用 adb install -r 覆盖安装成功，应用仍为 1.0.0 (10000)，5330 条书源、两本书架图书和阅读记录均保留。")
    add_bullet(document, "首轮搜索暴露“继续检测下一批”按钮事件未生效；改为搜索页原生按钮并补充回归断言后，真机点击 1.5 秒内进入探测状态。")
    add_bullet(document, "第二批报告有结果 1、空结果 2、失败 17；打开后识别《斗破苍穹》、作者九支书竹、目录 1642 章，第一章正文成功加载并超过 50 个字符。")
    add_bullet(document, "已知问题：搜索卡片先显示“未命名小说”，正文顶部残留 chap_tp(); theme();，下一轮继续修复标题提取与正文脚本清洗。")
    add_bullet(document, "玄幻聚合分类会跨多个提供者回退并稳定结束；移除 tcp:8765 映射后 12 秒内仍出现本地结果，测试完成后已恢复映射。")
    add_bullet(document, "五种 AI 音色均完成真实合成和手机播放，真实性刷新与缓存复跑通过；完整验收在 8/11 项后人工停止长时间连续播放。")
    add_bullet(document, "热修最终 APK 为 6,433,142 字节，SHA-256 为 D9FD511D04340AD0065A726BAEE4EDBD234467A4BCA155E000FEBC5BD1D9E1B8；v1/v2/v3 签名通过并覆盖安装成功。")
    return True


def append_stage7_section(document):
    if any(paragraph.text.strip() == STAGE7_SECTION_TITLE for paragraph in document.paragraphs):
        return False

    heading = document.add_heading(STAGE7_SECTION_TITLE, level=1)
    for run in heading.runs:
        set_east_asia_font(run)

    document.add_heading("16.1 搜索元数据与正文质量", level=2)
    add_bullet(document, "搜索解析不再提前写入“未命名小说”。结果分为 complete、needs_detail、invalid；只有标题和书籍 URL 均有效的结果可以直接展示。")
    add_bullet(document, "新增 hydrateSourceSearchResults()：每个来源最多补齐 3 个缺失标题结果、并发 2、详情超时 4 秒。补齐失败的项目不显示虚假卡片，并记录 SEARCH_RESULT_INCOMPLETE 或 DETAIL_METADATA_EMPTY。")
    add_bullet(document, "搜索与发现共用元数据补齐管线；只有至少一个完整结果时才把 runtimeV2.search 写为 passed，同书备用来源也必须拥有完整标题与 URL。")
    add_bullet(document, "新增独立正文清洗器与质量评估器：删除 script/style/noscript/template/svg/canvas、执行书源替换、转换 HTML、清理独立 JavaScript 调用行并合并重复片段。")
    add_bullet(document, "chapterCacheMeta 增加 sanitizerVersion、rawChars、cleanedChars。旧缓存首次读取时惰性重洗并原位升级，不清空章节，不改变书架 ID、章节索引或阅读进度。")
    add_bullet(document, "清洗后为空或以页面脚本为主时分别返回 CONTENT_EMPTY、CONTENT_NOISE；不足 50 字的合法短章仍可阅读，但不计入完整阅读通过率。")

    document.add_heading("16.2 候选调度与高频规则差异", level=2)
    add_bullet(document, "搜索与 Wi-Fi 预热按基础域名分散抽样，同一域名每轮最多 2 个来源；两个独立时间窗口均完整通过的来源进入最高优先级。")
    add_bullet(document, "SEARCH_RESULT_INCOMPLETE 与 CONTENT_NOISE 分别进入搜索、正文阶段冷却，不影响其他阶段；正文质量通过后才写入 runtimeV2.content=passed。")
    add_bullet(document, "声明式引擎补齐可复现的 CSS 类选择、简单 XPath、JSONPath 递归属性和 JSON 模板变量差异。YCK 代表样本 6645、6931 已由解析为空修复为完整阅读通过。")
    add_bullet(document, "任意 eval、Function、Java 类、文件系统、全局 DOM、验证码绕过与付费内容访问仍被拒绝；需要越界宿主能力的样本继续返回稳定安全错误。")

    document.add_heading("16.3 当前公开候选集首窗口", level=2)
    add_body(document, "基准报告升级为 schema v3，只保存 ID、配置哈希、阶段状态、耗时、错误码与内容长度统计，不保存正文、Cookie、Token、完整响应或完整书源配置。")
    rows = [
        ("当前候选配置", "120", "近期/中段/较早确定性分层，同域名最多 2 个"),
        ("合法且可导入", "120 / 120", "导入率 100%"),
        ("严格运行时分母", "33", "已达到至少 20 个"),
        ("完整阅读通过", "9 / 33", "27.27%，未达到至少 80%"),
        ("外部不可达或受限", "87", "不进入严格分母"),
        ("元数据失败", "1", "SEARCH_RESULT_INCOMPLETE"),
        ("合格正文抽样", "27 段", "噪声 0、短章 0"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(["验收项", "首窗口结果", "门槛与结论"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)
    add_bullet(document, "外部排除主要为 HTTP_BLOCKED 34、NETWORK_ERROR 25、TIMEOUT 13、HTTP_NOT_FOUND 11、HTTP_SERVER_ERROR 3、CAPTCHA_REQUIRED 1。")
    add_bullet(document, "严格分母内主要差异为 PARSE_EMPTY 19，另有 TOC_EMPTY、CONTENT_EMPTY、SEARCH_RESULT_INCOMPLETE、SEARCH_EMPTY、REQUEST_TEMPLATE_UNSUPPORTED 各 1。")
    warning = document.add_paragraph()
    warning_run = warning.add_run("验收结论：首窗口分母已达到 33，但完整阅读通过率仅 27.27%。第二窗口必须间隔至少 24 小时，当前尚未执行；PR #1 继续保持草稿，不能宣称全部或绝大 YCK 来源已稳定可读。")
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    warning_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    set_east_asia_font(warning_run)

    document.add_heading("16.4 自动测试、APK 与真机验收", level=2)
    add_bullet(document, "前端全量 106 / 106 passed；后端 SQLite 125 / 125 passed；H5 生产构建成功；git diff --check 通过。")
    add_bullet(document, "最终 APK 为 release/android-v2/V2.apk，1,343,966 字节，SHA-256 为 BF1703548EC296AEC9DEC381C1ADAF8A6AED69701A401F84BBDFCAA1BB8D1CEA；v1、v2、v3 签名均通过。")
    add_bullet(document, "REA-AN00 使用 adb install -r 覆盖安装成功，保留 5330 个来源、已有书架和阅读记录。移除 tcp:8765 映射后搜索“斗破苍穹”，20 秒内得到真实书名结果，未出现“未命名小说”。")
    add_bullet(document, "无后端状态下打开结果并识别 1642 章，首章正文成功加载；页面中 chap_tp、theme( 与占位标题均未出现。")
    add_bullet(document, "恢复 tcp:8765 后 FastAPI 健康检查正常；loli、uncle、youth、shota、recital 五种 AI 声音均已验证且可用。本阶段只做轻量回归，没有重复长时间三章播放。")

    document.add_heading("16.5 发布边界与下一步", level=2)
    add_number(document, "首窗口至少 24 小时后，用相同锁定清单和配置 SHA-256 执行第二窗口，并用 combine_source_acceptance_windows.mjs 生成双窗口结论。")
    add_number(document, "继续针对 PARSE_EMPTY 中占比最高且可复现的声明式差异增加脱敏夹具；每项兼容能力同时包含成功样本、越界拒绝和执行预算测试。")
    add_number(document, "只有两个窗口都达到分母至少 20、完整通过率至少 80%，且 PostgreSQL 16 CI、真机闭环与全部自动测试通过后，才把 PR #1 从草稿改为可审阅。")
    add_number(document, "第三方永久失效、主动屏蔽、登录、验证码和付费来源继续准确标记限制，不绕过访问控制；Android 阅读保持本地优先，后端只承载账号、同步、云 TTS 与 H5 代理。")
    return True


def append_stage7_replay_section(document):
    if any(paragraph.text.strip() == STAGE7_REPLAY_TITLE for paragraph in document.paragraphs):
        return False
    heading = document.add_heading(STAGE7_REPLAY_TITLE, level=2)
    for run in heading.runs:
        set_east_asia_font(run)
    add_bullet(document, "对首窗口 PARSE_EMPTY 样本逐项复测，新增标题链接 URL 推导、book.kind/chapter 阶段上下文、详情页目录回退、相对 URL 规则结果和换行分隔的受控 JS 声明语句。")
    add_bullet(document, "YCK 6808 修复为完整通过：搜索与详情成功，目录 1663 章，首章清洗后 3505 字符。")
    add_bullet(document, "YCK 6305 修复为完整通过：book.kind 正确传递到详情、目录和正文请求，多行受控 JS 模板在既有预算内运行，目录 1663 章，首章清洗后 2852 字符。")
    add_bullet(document, "连同此前修复的 6645、6931，首窗口已有 4 个 PARSE_EMPTY 样本完成搜索、详情、目录、正文闭环；前端全量回归为 108 / 108 passed。")
    add_bullet(document, "eval、Function、Java 类、文件系统、全局 DOM、循环与动态模块仍被拒绝。本次重放不修改首窗口原始统计，也不替代间隔至少 24 小时的第二窗口。")
    return True


def append_stage7_replay2_section(document):
    if any(paragraph.text.strip() == STAGE7_REPLAY2_TITLE for paragraph in document.paragraphs):
        return False
    heading = document.add_heading(STAGE7_REPLAY2_TITLE, level=2)
    for run in heading.runs:
        set_east_asia_font(run)
    add_bullet(document, "YCK 6247 完整通过：目录 50 章，首章清洗后 141 字符，验证复合类选择器兼容能力。")
    add_bullet(document, "YCK 6311 的 @onclick@js:result.match(...)[1] 规则通过只读属性、正则 match() 和安全数组索引实现，目录 1649 章，首章清洗后 3008 字符。")
    add_bullet(document, "首窗口共有 6 个原 PARSE_EMPTY 样本完成搜索、详情、目录和正文闭环；脱敏 Markdown 与 schema v3 JSON 同步保存 ID、配置哈希、阶段结果及长度统计。")
    add_bullet(document, "6238、5915、5813 的当前页面已不再包含配置声明的结果容器，5998 当前接口仍返回“请输入搜索词”；这些归类为站点或配置变化，不增加站点专用绕过逻辑。")
    add_bullet(document, "本批兼容仍受安全解释器预算约束，不开放 eval、Function、Java 类、文件系统或验证码绕过；重放结果不修改首窗口原始统计。")
    add_bullet(document, "前端 108 / 108、后端 SQLite 125 / 125 通过，生产 H5 与 APK 构建成功。最新 APK 为 1,348,062 字节，SHA-256 为 5F9E9658B559E0857A2E627365AE0C8624D71ACDA5C22883DEE21A17A5FCAAF2，v1/v2/v3 签名通过。")
    add_bullet(document, "REA-AN00 使用 adb install -r 保留数据覆盖安装成功，版本保持 1.0.0 (10000)，首次安装时间仍为 2026-05-29；第二窗口最早在北京时间 2026-08-14 15:47 后执行。")
    return True


def append_stage8_section(document):
    title = "17. 第八阶段全链路流畅度、内存与页面切换优化（2026-08-13）"
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        return False

    document.add_page_break()
    heading = document.add_heading(title, level=1)
    for run in heading.runs:
        set_east_asia_font(run)

    document.add_heading("17.1 导航、快照与后台任务", level=2)
    add_bullet(document, "底部标签移除固定 160ms 延迟，同一事件周期执行 switchTab；Android 自动性能档默认采用轻量模式，页面转场缩短为 80–100ms。")
    add_bullet(document, "5330 个书源改用版本化内存快照、ID 索引、30 条轻量分页和 120ms 筛选防抖；页面响应式数据不再保存全部 raw 配置。")
    add_bullet(document, "Android 原生存储桥新增分片批量读取；轻量书源索引和发现目录按数据修订号持久缓存，配置变化后精确失效。")
    add_bullet(document, "搜索、发现和正文运行状态先写内存并在 250ms 内合并落盘；应用进入后台时强制刷新。同步、索引和 Wi-Fi 预热在首屏后错峰执行。")

    document.add_heading("17.2 阅读、声音与性能诊断", level=2)
    add_bullet(document, "阅读器在正文分页完成后再延迟预加载下一章，并记录 reader.chapter.render；AI 声音列表和服务状态缓存 5 分钟，用户主动刷新时才强制重取。")
    add_bullet(document, "新增本地性能记录器和 Android PSS 采样，只保存阶段、耗时、数量和内存，不保存书名、正文、Cookie、Token 或完整书源。")
    add_bullet(document, "我的页面提供自动、流畅、完整三档性能模式，并在调试面板显示性能摘要、复制报告和清空报告。")

    document.add_heading("17.3 REA-AN00 真机结果", level=2)
    rows = [
        ("标签导航 P95", "4.9ms", "34 次应用内记录，低于 200ms 门槛"),
        ("30 次切换帧耗时", "P95 30ms / P99 42ms", "690 帧，卡顿帧 2.61%"),
        ("稳定态 PSS", "155,437KB", "相对约 269MB 基线下降约 42%"),
        ("连续切换内存", "约 +1.1MB", "稳定采样未持续增长，低于 20MB 门槛"),
        ("书源筛选分页", "P95 16.2ms", "5330 源、页面仅保留 30 条轻量行"),
        ("暖启动", "P95 177ms", "5 次 WARM 样本，低于 200ms 门槛"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(["指标", "结果", "结论"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)
    add_body(document, "说明：首次升级需在后台生成一次轻量索引和发现目录；生成后按修订号复用。采样期间观察到 WebView GC 前短时高 PSS，稳定态回落至约 155MB，报告按稳定态与连续增长口径验收。")

    document.add_heading("17.4 自动验证、交付边界与后续", level=2)
    add_bullet(document, "前端 111 / 111 passed；后端 SQLite 125 / 125 passed；生产 H5 和 Android APK 构建成功。原生书源分片改为每批最多 16 个读取后，完整 5330 源首次加载采样峰值由 391,717KB 降至 212,276KB，约 0.9 秒回落到 151,123KB。最终 APK 为 1,352,158 字节，SHA-256 为 A47BE2B0057D94D391ED314FF96446E4FC72CD41A03364269964609BF9034CC5，v1/v2/v3 签名通过。")
    add_bullet(document, "覆盖安装保留 5330 个书源、书架、章节缓存、阅读进度和首次安装时间；不新增运行时依赖或数据库表。")
    add_bullet(document, "阶段七第二时间窗口最早为北京时间 2026-08-14 15:47，本轮未提前执行或修改冻结清单；首窗口仍为 9/33（27.27%），PR #1 继续保持 Draft。")
    add_number(document, "第二时间窗口到期后用原锁定清单复测并生成双窗口合并报告；未达到分母不少于 20、完整通过率不少于 80% 前不宣称绝大书源稳定可读。")
    add_number(document, "继续对首次升级索引构建、发现首屏和阅读章节渲染做真机性能采样；对稳定复现的慢阶段实施最小化优化。")
    return True


def append_stage9_section(document):
    title = "18. 第九阶段发现目录预生成与首次进入优化（2026-08-13）"
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        return False

    document.add_heading(title, level=1)
    document.add_heading("18.1 实现与生命周期", level=2)
    add_bullet(document, "书源索引完成后在浏览器空闲回调中预生成发现候选、入口和分类目录；用户第一次进入发现页时优先读取轻量缓存。")
    add_bullet(document, "配置导入、覆盖、启停或编辑后按书源修订号失效；应用退后台时取消未执行任务，回到前台后按有效快照重新安排。")
    add_bullet(document, "预生成只整理本地元数据，不访问第三方站点，不保存正文、Cookie、Token 或完整响应，也不增加运行时依赖。")

    document.add_heading("18.2 自动验证与 REA-AN00 真机结果", level=2)
    rows = [
        ("前端测试", "112 / 112 passed", "新增预生成、取消和失效回归"),
        ("发现页暖切换", "500ms 首屏完整", "48 个发现源、12 个分类可见"),
        ("GC 后稳态 PSS", "142,817KB", "覆盖安装后稳定采样"),
        ("数据保留", "5330 源 / 3 本书架", "首次安装时间保持 2026-05-29"),
        ("最终 APK", "1,356,254 字节", "v1/v2/v3 签名通过"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(["指标", "结果", "结论"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)
    add_body(document, "阶段九 APK SHA-256：5DD2B5DE7D79D4498BBCF60B0B10A0072F14B7FA6DA54A72BE58AA7D08E33575。")

    document.add_heading("18.3 发布边界与下一步", level=2)
    add_bullet(document, "阶段七第二时间窗口最早为北京时间 2026-08-14 15:47；当前仍未执行，首窗口 9/33（27.27%）保持冻结，PR #1 继续保持 Draft。")
    add_number(document, "按冻结清单执行第二窗口并生成双窗口报告；未达到 80% 时修复前三类通用规则差异。")
    add_number(document, "达到门槛后完成关闭后端阅读、恢复后端同步和五种 AI 声音回归，再将 PR 转为可审阅。")
    return True


def main():
    document = Document(DOCX_PATH)
    if any(paragraph.text.strip() == SECTION_TITLE for paragraph in document.paragraphs):
        stage4_updated = False
        stage6_updated = False
        stage8_updated = False
        for paragraph in document.paragraphs:
            if "不自动透传 Cookie、Authorization 或原始请求体" in paragraph.text:
                paragraph.text = "兼容相对路径 POST 搜索与首页跨域跳转：先以无敏感信息 GET 确认最终站点，再向新源站重建只包含搜索参数的 POST；不跨域透传 Cookie、Authorization 或 Proxy-Authorization。"
                for run in paragraph.runs:
                    set_east_asia_font(run)
                stage4_updated = True
            if "代码已通过自动测试，尚待新 APK 复测" in paragraph.text:
                paragraph.text = "真机搜索斗破苍穹在 20 秒内探测 20 个来源但未得到结果；随后增加真实通过源冷启动排序并修复无结果页入口，代码已通过自动测试并进入最终 APK，尚待设备重新连接后复测。"
                for run in paragraph.runs:
                    set_east_asia_font(run)
                stage6_updated = True
            if "最终 H5/APK 重建被外部执行审批层拒绝" in paragraph.text:
                paragraph.text = "2026-08-13 最终 H5/APK 重建成功；APK 为 6,260,746 字节，SHA-256 为 AEBDCC9D87F7E4CA2A3B022DBAC20371848D83B1F94B0295867727D6E5B3405A，v1/v2/v3 签名通过。无线 ADB 在覆盖安装前断开，最终包尚未完成 REA-AN00 复测。"
                for run in paragraph.runs:
                    set_east_asia_font(run)
                stage6_updated = True
            if "权限恢复后重建 H5/APK" in paragraph.text:
                paragraph.text = "设备重新连接后覆盖安装最终 APK，并复测搜索、发现、备用线路、断网缓存与 AI 语音试听。"
                for run in paragraph.runs:
                    set_east_asia_font(run)
                stage6_updated = True
            if "9AAE5CC3AEE128ABC9A5477EE79E768533C7A1CA1F6308A2F814028A58682C16" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "9AAE5CC3AEE128ABC9A5477EE79E768533C7A1CA1F6308A2F814028A58682C16",
                    "3C8BA0EFD6BCFB781188CA28063612CB74A1EFEACAA64B7D23D0DE656B542967"
                ).replace("，v1/v2/v3 签名通过。", "，v1/v2/v3 签名通过；最终包覆盖安装并重启后仍显示 5330 源。")
                for run in paragraph.runs:
                    set_east_asia_font(run)
                stage4_updated = True
            if paragraph.text.startswith(("前端 110 / 110 passed；后端 SQLite 125 / 125 passed", "前端 111 / 111 passed；后端 SQLite 125 / 125 passed")) and "A47BE2B0" not in paragraph.text:
                paragraph.text = "前端 111 / 111 passed；后端 SQLite 125 / 125 passed；生产 H5 和 Android APK 构建成功。原生书源分片改为每批最多 16 个读取后，完整 5330 源首次加载采样峰值由 391,717KB 降至 212,276KB，约 0.9 秒回落到 151,123KB。最终 APK 为 1,352,158 字节，SHA-256 为 A47BE2B0057D94D391ED314FF96446E4FC72CD41A03364269964609BF9034CC5，v1/v2/v3 签名通过。"
                for run in paragraph.runs:
                    set_east_asia_font(run)
                stage8_updated = True
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() == "110 / 110 passed":
                        cell.text = "111 / 111 passed"
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                set_east_asia_font(run)
                        stage8_updated = True
        stage2_added = append_stage2_section(document)
        stage3_added = append_stage3_section(document)
        stage4_added = append_stage4_section(document)
        stage5_added = append_stage5_section(document)
        stage6_added = append_stage6_section(document)
        stage6_final_added = append_stage6_final_section(document)
        stage7_added = append_stage7_section(document)
        stage7_replay_added = append_stage7_replay_section(document)
        stage7_replay2_added = append_stage7_replay2_section(document)
        stage8_added = append_stage8_section(document)
        stage9_added = append_stage9_section(document)
        if stage4_updated or stage6_updated or stage8_updated or stage2_added or stage3_added or stage4_added or stage5_added or stage6_added or stage6_final_added or stage7_added or stage7_replay_added or stage7_replay2_added or stage8_added or stage9_added:
            saved_path = save_document(document)
            print(f"Updated: {saved_path}")
        else:
            print(f"Sections already present: {DOCX_PATH}")
        return
    else:
        document.add_page_break()
        heading = document.add_heading(SECTION_TITLE, level=1)
        for run in heading.runs:
            set_east_asia_font(run)

        document.add_heading("10.1 架构决策", level=2)
    add_bullet(document, "Android APK 采用本地优先：URL、文件、二维码、深链和 3.x JSON 在手机本地识别、预览、去重、保存与运行。")
    add_bullet(document, "新增 NovelReaderHttp 原生桥；APK 的外部书源请求固定为“原生 HTTP → 必要时 WebView 渲染”，不再默认访问 localhost:8765。")
    add_bullet(document, "后端改为可选云服务，仅承担账号同步、云书架、云 TTS、H5 跨域代理和基础来源兼容。Android 联网阅读不要求连接电脑后端。")
    add_bullet(document, "H5 无后端时只保证同源或允许 CORS 的来源；这是浏览器平台限制，不在 UI 中伪装成书源失效。")

    document.add_heading("10.2 统一导入与数据迁移", level=2)
    add_bullet(document, "新增 resolveSourceImport、previewSourceImport、applySourceImport、requestSourceText、runSourceReadingFlow 公共接口。")
    add_bullet(document, "支持对象/数组/sources 包装、BOM、JSON/TXT、剪贴板、二维码、yuedu://、legado://、booksource://、JSON URL 和 YCK 详情页。")
    add_bullet(document, "新增稳定 sourceKey（规范化名称 + 基础 URL），本地存储迁移到版本 3；旧 id 原样保留，避免书架引用断裂。")
    add_bullet(document, "合法但受限的来源保存为禁用状态；只有无效 JSON 或缺少名称/基础 URL 时拒绝。状态统一为 ready、partial、needs_login、blocked、invalid。")

    document.add_heading("10.3 Android 传输与 3.x 兼容矩阵", level=2)
    add_body(document, "NovelReaderHttp 支持 GET/POST、请求头/body、UTF-8/GBK/GB2312、Cookie 隔离、Referer、User-Agent、重定向、超时、限速、并发和响应体限制。私网、回环、file:、content: 等目标默认禁止；跨域重定向会清除 Cookie/Authorization。")
    rows = [
        ("CSS / 属性 / 索引", "支持", "支持", "基础支持", "含 text/html/href/src/textNodes/ownText"),
        ("XPath", "支持", "浏览器支持", "暂不支持", "Android WebView DOM 执行"),
        ("JSONPath / 正则", "支持", "支持", "基础支持", "含 || 回退与 && 拼接"),
        ("GET/POST/headers/charset", "支持", "受 CORS 限制", "支持", "APK 走原生桥"),
        ("目录/正文多页", "支持", "支持", "暂未扩展", "默认最多 5 页"),
        ("安全 JS 子集", "支持", "支持", "不执行", "字符串/数组/JSON/URL/Base64 + 执行预算"),
        ("任意 java.* / eval", "阻止", "阻止", "阻止", "保存为禁用，不执行"),
        ("登录/验证码/付费", "人工处理", "受限", "不绕过", "只显示限制原因"),
    ]
    table = document.add_table(rows=1, cols=5)
    for index, value in enumerate(["能力", "Android", "H5 无后端", "后端", "说明"]):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table)

    document.add_heading("10.4 后端契约", level=2)
    add_bullet(document, "新增鉴权接口 POST /api/sources/import/preview，只分析书源文本，不写数据库。")
    add_bullet(document, "POST /api/sources/import 增加 source_url、import_method、duplicate_strategy，并返回 updated/skipped/unsupported 计数和逐项平台能力。")
    add_bullet(document, "不新增数据库表；沿用 raw_json、compatibility、health_status 和加密会话字段，保持用户隔离、软删除恢复与事务一致性。")

    document.add_heading("10.5 真实 YCK 基准与结论", level=2)
    add_body(document, "抓取页：1、28、56；按近期/中段/较早三层固定 200 个合法文字源。报告只保存 ID、抓取时间、SHA-256、阶段状态、耗时和错误码，不保存正文、Cookie、Token 或完整书源 JSON。")
    add_bullet(document, "有效文字 JSON 200；可导入 200；导入率 100%，达到 ≥95% 门槛。")
    add_bullet(document, "静态状态：ready 86、partial 32、needs_login 10、blocked 72。")
    add_bullet(document, "严格合格候选 38；桌面真实完整流程 0/38：SEARCH_FAILED 30、SEARCH_EMPTY 6、TIMEOUT 2。")
    add_bullet(document, "代表源 7163、7298 单独探测均完成搜索、详情、目录和正文，目录分别为 1663、999 章。")
    warning = document.add_paragraph()
    warning_run = warning.add_run("验收结论：随机基准的 ≥80% 完整阅读目标未通过，当前版本不能宣称“绝大书源可读”。")
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    warning_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    set_east_asia_font(warning_run)

    document.add_heading("10.6 验证证据与发布阻断项", level=2)
    add_bullet(document, "前端 92 个 *.test.mjs 文件执行；后端 SQLite 全量 125 passed；PostgreSQL 16 由 GitHub Actions 继续验证。")
    add_bullet(document, "H5 生产构建及本地导入持久化验收通过；APK release/android-v2/V2.apk 为 1,496,142 字节，SHA-256 为 11D11EC281FBCF93D43B3B9A34C9900365238CAB93A523E92ABECBF8CBF39BCD，v1/v2/v3 签名通过。")
    add_bullet(document, "本轮没有可用 Android 真机/ADB 设备；关闭 8765 后扫码导入、覆盖安装、重启续读、断网缓存和内置摄像头扫码仍是发布阻断项。")
    add_number(document, "以 SEARCH_FAILED 样本补齐高频 3.x 请求脚本和受控宿主 API 映射，每次只扩展白名单能力。")
    add_number(document, "进一步区分站点失效、规则不支持、网络超时与无关键词结果；完整阅读率达到 ≥80% 前不合并正式发布。")
    add_number(document, "在无电脑后端的 Android 真机完成 URL、文件、二维码、深链同源去重和完整阅读闭环，并保存录屏与 APK 哈希。")

    append_stage2_section(document)
    append_stage3_section(document)
    append_stage4_section(document)
    append_stage5_section(document)
    append_stage6_section(document)
    append_stage6_final_section(document)
    append_stage7_section(document)
    append_stage7_replay_section(document)
    append_stage7_replay2_section(document)
    append_stage8_section(document)
    append_stage9_section(document)
    saved_path = save_document(document)
    print(f"Updated: {saved_path}")


if __name__ == "__main__":
    main()
