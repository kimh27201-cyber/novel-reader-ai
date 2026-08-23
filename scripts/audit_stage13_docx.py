import json
import sys
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else ROOT / "docs" / "解码阅读-V3阶段开发记录-2026-08-11.docx"
STAGE_TITLE = "22. 第十三阶段真实可读率收口与 V3.0 发布门禁（2026-08-24）"
REQUIRED_ENTRIES = {
    "[Content_Types].xml",
    "_rels/.rels",
    "word/document.xml",
    "word/styles.xml",
    "word/_rels/document.xml.rels",
}
REQUIRED_MARKERS = {
    "stage_title": STAGE_TITLE,
    "gate_rate": "52.73%",
    "gate_status": "PREQUALIFICATION_GATE_FAILED",
    "frontend_tests": "120 / 120 passed",
    "source_count": "5330 源",
    "apk_hash": "B3EC611D24BB1814469C65D2724E176EF1A8DDF3549D9FA46831FAA529DBD147",
}


with zipfile.ZipFile(DOCX_PATH) as archive:
    entries = set(archive.namelist())
    corrupt_entry = archive.testzip()
    missing_entries = sorted(REQUIRED_ENTRIES - entries)
    relationship_count = archive.read("word/_rels/document.xml.rels").count(b"<Relationship")

document = Document(DOCX_PATH)
paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
all_text = "\n".join(paragraphs + [cell.text for table in document.tables for row in table.rows for cell in row.cells])
headings = [
    paragraph.text.strip()
    for paragraph in document.paragraphs
    if paragraph.text.strip() and str(paragraph.style.name or "").lower().startswith("heading")
]
marker_status = {name: value in all_text for name, value in REQUIRED_MARKERS.items()}
stage_occurrences = paragraphs.count(STAGE_TITLE)

result = {
    "path": str(DOCX_PATH),
    "bytes": DOCX_PATH.stat().st_size,
    "zipCorruptEntry": corrupt_entry or "",
    "missingRequiredEntries": missing_entries,
    "relationshipCount": relationship_count,
    "paragraphCount": len(document.paragraphs),
    "tableCount": len(document.tables),
    "headingCount": len(headings),
    "stage13TitleOccurrences": stage_occurrences,
    "markers": marker_status,
    "passed": not corrupt_entry
    and not missing_entries
    and stage_occurrences == 1
    and all(marker_status.values()),
}

print(json.dumps(result, ensure_ascii=False, indent=2))
raise SystemExit(0 if result["passed"] else 1)
